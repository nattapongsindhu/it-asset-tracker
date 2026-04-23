from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "IT-Asset-Tracker-Portfolio-Lite-Plan-TH.md"
OUTPUT = ROOT / "docs" / "IT-Asset-Tracker-Portfolio-Lite-Plan-TH.docx"
THAI_FONT = "Leelawadee UI"


def set_style_font(style_name: str, font_name: str, size: int | None = None) -> None:
    style = DOC.styles[style_name]
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    r_fonts = style._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def add_paragraph_with_font(text: str, style: str | None = None, center: bool = False) -> None:
    paragraph = DOC.add_paragraph(style=style)
    if center:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(text)
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), THAI_FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


DOC = Document()

set_style_font("Normal", THAI_FONT, 11)
set_style_font("Heading 1", THAI_FONT, 16)
set_style_font("Heading 2", THAI_FONT, 13)
set_style_font("Heading 3", THAI_FONT, 11)


def main() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    first_title = True
    for raw_line in lines:
        line = raw_line.rstrip()

        if not line:
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if first_title:
                add_paragraph_with_font(text, center=True)
                DOC.paragraphs[-1].runs[0].bold = True
                DOC.paragraphs[-1].runs[0].font.size = Pt(20)
                first_title = False
            else:
                add_paragraph_with_font(text)
                DOC.paragraphs[-1].style = DOC.styles["Heading 1"]
            continue

        if line.startswith("## "):
            add_paragraph_with_font(line[3:].strip())
            DOC.paragraphs[-1].style = DOC.styles["Heading 1"]
            continue

        if line.startswith("### "):
            add_paragraph_with_font(line[4:].strip())
            DOC.paragraphs[-1].style = DOC.styles["Heading 2"]
            continue

        if line.startswith("- "):
            add_paragraph_with_font(line[2:].strip(), style="List Bullet")
            continue

        add_paragraph_with_font(line)

    DOC.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
